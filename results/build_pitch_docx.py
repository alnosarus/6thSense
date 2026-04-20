"""Generate 6thSense VC pitch DOCX mirroring the typst version."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


GREY_TEXT = RGBColor(0x55, 0x55, 0x55)
GREY_LIGHT = RGBColor(0x88, 0x88, 0x88)
CALLOUT_FILL = "F2F2F2"
TABLE_BORDER = "B4B4B4"


def shade_cell(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_table_borders(table, color_hex=TABLE_BORDER):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tbl_pr.append(borders)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    runs_from_markup(p, text, base_size=10.5)
    set_table_borders(table, "DDDDDD")
    doc.add_paragraph()


def runs_from_markup(paragraph, text, base_size=11):
    """Minimal markup: *bold* segments, rest regular."""
    parts = text.split("*")
    bold = False
    for part in parts:
        if part:
            run = paragraph.add_run(part)
            run.font.size = Pt(base_size)
            run.bold = bold
        bold = not bold


def add_heading(doc, text, level):
    sizes = {1: 20, 2: 14, 3: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes[level])
    if level == 3:
        run.italic = True
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    runs_from_markup(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        runs_from_markup(p, item)


def build_table(doc, headers, rows, col_widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Inches(w)

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        shade_cell(cell, "EAEAEA")

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            runs_from_markup(p, value, base_size=10)

    set_table_borders(table)
    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # page + default style
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── cover block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("6thSense")
    run.bold = True
    run.font.size = Pt(28)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("Robots have five senses. We build the sixth.")
    run.font.size = Pt(12)
    run.font.color.rgb = GREY_TEXT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "VC Pitch Brief  ·  Competition  ·  Why We Win  ·  Why We Are The Team"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GREY_LIGHT

    doc.add_paragraph()

    add_body(
        doc,
        "*Positioning.* We sell the *dataset*. The rig is how we own the cost "
        "curve. 6thSense ships model-ready, multimodal tactile + egocentric "
        "episodes for contact-rich robot learning — captured on a rig we "
        "engineer in-house at ~5× lower cost than the incumbent.",
    )

    # ── 1. Competition
    add_heading(doc, "Competition", 1)
    add_body(
        doc,
        "Robot-learning teams trying to solve contact-rich manipulation today "
        "have five procurement options. None of them deliver *packaged, "
        "model-ready tactile + ego episodes* at a price that makes training "
        "economics work.",
    )

    build_table(
        doc,
        headers=["Category", "Representative players", "Why they don't solve the real problem"],
        rows=[
            [
                "Tactile glove incumbents",
                "HaptX (~$5k/glove), Manus, SenseGlove",
                "Hardware-only. 5× our unit cost. No sync, no QC, no labels, no dataset schema.",
            ],
            [
                "Fingertip tactile sensors",
                "GelSight, DIGIT (Meta), Xela, Contactile",
                "Point-contact only. No full-hand skin coverage. Not a data product — a sensor SKU.",
            ],
            [
                "Egocentric video data",
                "Ego4D, Project Aria (Meta), Ego-Exo4D",
                "No touch modality. Academic license. Not scoped to robot training schemas.",
            ],
            [
                "Vertically-integrated robotics co's",
                "Physical Intelligence, 1X, Figure, Dyna, Skild",
                "Running in-house data ops — *badly and expensively*. These are *future customers*, not rivals.",
            ],
            [
                "Generalist data / labeling co's",
                "Scale AI, Surge, Invisible",
                "No hardware. No multimodal sync. No tactile domain expertise. Commodity labeling only.",
            ],
            [
                "Open benchmarks",
                "DROID, Open-X Embodiment, RH20T",
                "RGB + proprioception only. Complementary to us, not competitive — they become citations, not threats.",
            ],
        ],
        col_widths=[1.4, 1.9, 2.9],
    )

    add_callout(
        doc,
        "*Bottom line.* Nobody else is selling tactile + egocentric data as a "
        "managed service at our cost structure. HaptX is the nearest adjacent "
        "and they're a hardware company charging 5× more for something with "
        "worse downstream fit.",
    )

    # ── 2. Why We Win
    add_heading(doc, "Why We Win", 1)
    add_body(
        doc,
        "Three compounding advantages, each defensible on its own, reinforcing "
        "across the others.",
    )

    add_heading(doc, "1. A 5× cost advantage on the capture rig", 2)
    add_body(
        doc,
        "HaptX — the only serious incumbent in full-hand tactile — sells "
        "gloves at roughly *$5,000 per glove*. Our rig, built on JQ Industries "
        "e-skin + Intel RealSense, lands at *~$1,000 per glove* with "
        "line-of-sight lower.",
    )
    add_body(
        doc,
        "At our unit cost, *the same raised dollar fields 5× more gloves* — "
        "which compounds directly into:",
    )
    add_bullets(
        doc,
        [
            "5× the captured hours",
            "5× the scene coverage in the catalog",
            "5× the dataset moat per funding round",
        ],
    )
    add_body(
        doc,
        "The rig is our *factory*, not our product. We never ship it, never "
        "support it, never RMA it. It is internal capex amortized across "
        "every episode we sell.",
    )

    add_heading(doc, "2. Dataset ops as a product, not a byproduct", 2)
    add_body(
        doc,
        "Every robotics lab underestimates the *80%* of the work that sits "
        "between raw sensor dumps and a trainable episode:",
    )
    add_bullets(
        doc,
        [
            "Multi-rate sync: tactile ~1 kHz ↔ RGB-D 30–60 Hz ↔ IMU 200 Hz, on one clock",
            "Per-channel tactile calibration and drift handling",
            "Contact-phase segmentation and failure-mode taxonomy",
            "QC flagging with stated signal reliability boundaries",
            "Packaging to the format each customer's trainer expects",
        ],
    )
    add_body(
        doc,
        "This is an *operations moat*. It is earned through thousands of "
        "captured-then-QC'd hours and codified tooling — not a two-week "
        "script anyone can replicate.",
    )

    add_heading(doc, "3. Demand is already proven — pre-collection", 2)
    add_body(
        doc,
        "We are pre-collection, yet we are carrying *Letters of Intent* from:",
    )
    add_bullets(
        doc,
        [
            "*Anduril* (defense / autonomy)",
            "*Handshake* (commercial robotics)",
            "Additional smaller robotics companies",
            "University research labs",
        ],
    )
    add_body(
        doc,
        "Defense, commercial humanoid, and academic signal — *all three* — "
        "before episode one. This is the demand-proof every data-infrastructure "
        "VC looks for.",
    )

    add_callout(
        doc,
        "*Data-as-a-service multiples, hardware-company capital efficiency.* "
        "We price like a data company, compound like a data company, and "
        "build like a hardware company — so our cost base never catches up "
        "to us.",
    )

    # ── 3. Why We Are The Team
    add_heading(doc, "Why We Are The Team", 1)
    add_body(
        doc,
        "The problem has exactly three hard seats. We have exactly three "
        "founders mapped to them.",
    )

    build_table(
        doc,
        headers=["Founder", "Background", "Owns"],
        rows=[
            [
                "*Matt*",
                "Ex-Tesla hardware specialist",
                "Custom rig design, mechanical + electrical integration, BOM engineering that drives the $5k → $1k glove cost collapse and keeps driving it down.",
            ],
            [
                "*Ronak*",
                "Software / multimodal ML",
                "The insight + labeling pipeline: sync across tactile, RGB, and depth; calibration; contact-phase labeling; automated QC. The software moat that turns raw capture into trainable episodes.",
            ],
            [
                "*Alex*",
                "Founder, ML-readiness",
                "Product, customer programs, ML-signal validation. Translating fuzzy \"we need touch data\" asks into scoped, QC-able episode contracts. Won't sell a signal we can't prove is learnable.",
            ],
        ],
        col_widths=[1.0, 1.5, 3.7],
    )

    add_heading(doc, "Why this seat configuration matters", 2)
    add_body(
        doc,
        "A pure-software team *cannot build the rig*. A pure-hardware team "
        "*cannot ship model-ready episodes*. Everyone else trying to solve "
        "this problem is missing at least one seat — which is why they're "
        "either selling $5k gloves or paying internal engineers $300k/yr "
        "to run data ops.",
    )
    add_body(
        doc,
        "We are one of the very few teams with hardware IP, software IP, and "
        "ML-readiness judgment under one roof. That is the *entire* reason "
        "we can field LOIs from Anduril and Handshake while pre-collection.",
    )

    add_callout(
        doc,
        "*One line to close on.* We're a dataset company — we built the rig "
        "because no one sells a capture platform at a price that makes the "
        "dataset economics work, and we have the team to do both.",
    )

    out_path = "/Users/alexnoh/Desktop/6thsense/results/6thsense_vc_pitch.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
